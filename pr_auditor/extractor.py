"""Text extraction utilities for PDF, DOCX, and Excel files."""
import io


def extract_pdf_text(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    full_text = ""
    for page_num, page in enumerate(reader.pages, 1):
        extracted = page.extract_text() or ""
        full_text += f"\n[PAGE {page_num}]\n{extracted}"
    return full_text


def extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, filename: str) -> str:
    fn = filename.lower()
    if fn.endswith(".pdf"):
        return extract_pdf_text(file_bytes)
    elif fn.endswith((".docx", ".doc")):
        return extract_docx_text(file_bytes)
    return ""


def read_excel_as_text(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                rows_text.append(" | ".join(str(c) if c is not None else "" for c in row))
        if rows_text:
            parts.append(f"\n=== SHEET: {sheet_name} ===\n" + "\n".join(rows_text))
    wb.close()
    return "\n".join(parts)
